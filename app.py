import streamlit as st
import google.generativeai as genai
import csv
import io
import pandas as pd
import re
import random
from datetime import datetime, timedelta
from docx import Document

# ==========================================
# 1. ページ設定 & カスタムCSS
# ==========================================
st.set_page_config(
    page_title="ツナグ先生 - 統合校務支援システム (V11.0 日常指導＆仮名化セキュリティ強化版)",
    page_icon="🏫",
    layout="wide"
)

st.markdown("""
    <style>
    .block-container {
        padding-top: 3.5rem !important;
        padding-bottom: 1rem !important;
    }
    
    .student-card { 
        background-color: #f8f9fa; 
        border: 1px solid #dee2e6; 
        padding: 15px; 
        border-radius: 8px; 
        margin-bottom: 15px; 
    }

    .alert-card {
        background-color: #fff3cd;
        border: 1px solid #ffeeba;
        color: #856404;
        padding: 12px;
        border-radius: 6px;
        margin-bottom: 10px;
    }

    .stamp-btn button {
        width: 100%;
        margin-bottom: 5px;
    }
    </style>
""", unsafe_allow_html=True)

# ==========================================
# 2. ダミーデータ生成ロジック
# ==========================================

@st.cache_data
def generate_full_dummy_data():
    surnames = ["佐藤", "鈴木", "高橋", "田中", "伊藤", "渡辺", "山本", "中村", "小林", "加藤", 
                "吉田", "山田", "佐々木", "山口", "松本", "井上", "木村", "林", "斎藤", "清水"]
    male_names = ["蓮", "悠真", "湊", "大翔", "樹", "陽翔", "悠人", "颯太", "陸", "翔太"]
    female_names = ["葵", "陽葵", "凛", "結菜", "芽依", "詩", "結愛", "莉子", "咲良", "結衣"]

    classes_config = [
        ("1年1組", 40),
        ("1年5組", 40),
        ("2年2組", 40),
        ("3年3組", 40),
        ("3年5組", 40)
    ]

    master_list = []
    score_list = []
    
    random.seed(42)

    for cls_name, count in classes_config:
        for num in range(1, count + 1):
            gender = "男" if num % 2 != 0 else "女"
            sname = surnames[(num - 1) % len(surnames)]
            gname = male_names[(num - 1) % len(male_names)] if gender == "男" else female_names[(num - 1) % len(female_names)]
            full_name = f"{sname} {gname}"

            master_list.append({
                "クラス": cls_name,
                "出席番号": num,
                "氏名": full_name,
                "性別": gender,
                "ステータス": "在籍",
                "異動日": "",
                "備考": "担任クラス" if cls_name == "1年1組" else "授業担当"
            })

            is_absent = (num == 13 or num == 27)
            mid_score = None if is_absent else random.randint(45, 98)
            final_score = random.randint(50, 100)
            estimated_score = 65 if is_absent else None
            
            k1 = random.randint(55, 98)
            k2 = random.randint(50, 95)
            k3 = random.randint(60, 100)
            
            eval_1 = "A" if k1 >= 80 else ("B" if k1 >= 50 else "C")
            eval_2 = "A" if k2 >= 80 else ("B" if k2 >= 50 else "C")
            eval_3 = "A" if k3 >= 80 else ("B" if k3 >= 50 else "C")

            avg = (k1 + k2 + k3) / 3
            grade = 5 if avg >= 85 else (4 if avg >= 70 else (3 if avg >= 55 else (2 if avg >= 40 else 1)))

            score_list.append({
                "クラス": cls_name,
                "出席番号": num,
                "氏名": full_name,
                "中間テスト": mid_score,
                "期末テスト": final_score,
                "見込み点": estimated_score,
                "観点1_知識(点)": k1,
                "観点1_評価": eval_1,
                "観点2_思考(点)": k2,
                "観点2_評価": eval_2,
                "観点3_主体性(点)": k3,
                "観点3_評価": eval_3,
                "自動評定": grade,
                "★確定評定": grade,
                "調整フラグ": "―",
                "調整理由": "",
                "総合所見": f"{'課題に対して粘り強く思考し、工夫して解決策を導くことができました。' if grade >= 4 else '基礎的な計算力の定着が見られ、授業中の挙手・発言も意欲的です。'}"
            })

    logs_list = []
    base_date = datetime.now()
    memo_templates = [
        ("数学", "方程式の立式において、自力で関係性を見つけ出して正解を導き出すことができた。"),
        ("数学", "グループワークで解き方に悩んでいる班員に対して丁寧にやり方を教えていた。"),
        ("総合・行動", "行事の実行委員に立候補し、クラス全体の意見をスムーズに集約・発表した。"),
        ("総合・行動", "清掃活動において自分の担当場所が終わった後、進んで共有スペースの掃除を手伝った。"),
        ("特別活動", "朝の読書時間に毎日落ち着いて読書に取り組み、クラスの静寂な雰囲気作りに貢献した。"),
        ("国語・他", "朝の会での1分スピーチにて、自分の体験に基づいた説得力のある発表を行った。")
    ]

    c1_students = [m for m in master_list if m["クラス"] == "1年1組"]
    for idx, st_item in enumerate(c1_students):
        if st_item["出席番号"] % 2 != 0:
            days_ago = random.randint(1, 10)
        else:
            days_ago = random.randint(25, 40)
            
        d_str = (base_date - timedelta(days=days_ago)).strftime("%Y-%m-%d")
        cat, memo_text = memo_templates[idx % len(memo_templates)]
        logs_list.append({
            "日付": d_str,
            "クラス": "1年1組",
            "出席番号": st_item["出席番号"],
            "氏名": st_item["氏名"],
            "対象分野": cat,
            "観察メモ": memo_text
        })

    return pd.DataFrame(master_list), pd.DataFrame(score_list), pd.DataFrame(logs_list)


# ==========================================
# 3. セッション状態の初期化
# ==========================================

if "school_year" not in st.session_state:
    st.session_state.school_year = "2026"
if "teacher_name" not in st.session_state:
    st.session_state.teacher_name = "山田 太郎"

df_master, df_scores, df_logs = generate_full_dummy_data()

if "student_master" not in st.session_state:
    st.session_state.student_master = df_master

if "daily_logs" not in st.session_state:
    st.session_state.daily_logs = df_logs

if "subject_scores" not in st.session_state:
    st.session_state.subject_scores = df_scores

# APIキーの初期化（st.secretsがあれば取得）
api_key = ""
try:
    if "GEMINI_API_KEY" in st.secrets:
        api_key = st.secrets["GEMINI_API_KEY"]
except Exception:
    pass

# サイドバー全メニュー統合
menu_options = [
    "--- 🏠 学級担任機能 ---",
    "📝 ① 日々メモ・クイックスタンプ & 所見",
    "🔔 ② 学級日常ダッシュボード（観察アラート）",
    "📁 ③ CSV一括生成",
    "🔍 ④ 所見データ自動校正",
    "💬 ⑤ 蓄積連動カルテ",
    "🔄 ⑧ 担任用 全教科成績集約",
    "🖨️ ⑨ 完成版プレビュー ＆ ファイルダウンロード（通知表・指導要録）",
    "--- 📚 教科担当機能 ---",
    "📊 ⑥ 成績・観点A/B/C算出＆人間調整",
    "📈 ⑦ 学期推移ダッシュボード",
    "⚙️ ⓪ 担任＆授業担当 名簿管理"
]

if "selected_menu" not in st.session_state:
    st.session_state.selected_menu = "📝 ① 日々メモ・クイックスタンプ & 所見"

with st.sidebar:
    st.title("🏫 ツナグ先生")
    st.caption(f"統合校務支援システム (V11.0 / {st.session_state.school_year}年度)")
    st.markdown("---")

    st.markdown("**📌 機能メニュー選択**")
    nav_selection = st.radio("機能選択", options=menu_options, index=1, label_visibility="collapsed")
    
    if not nav_selection.startswith("---"):
        st.session_state.selected_menu = nav_selection

    st.markdown("---")
    st.markdown("**⚙️ システム基本設定 & セキュリティ**")
    st.session_state.school_year = st.text_input("年度設定", st.session_state.school_year)
    st.session_state.teacher_name = st.text_input("担任教員名", st.session_state.teacher_name)
    
    # 個人情報匿名化プロテクトの切り替えスイッチ
    use_anonymize = st.toggle("🔒 AI送信時の個人情報仮名化プロテクト", value=True, help="AIプロンプト送信時に氏名を『生徒A』等に自動変換し、個人情報の漏洩を強力に防止します。")

    # 🔑 Gemini APIキー入力エリア（追加・更新部分）
    st.markdown("---")
    st.markdown("**🔑 Google Gemini APIキー設定**")
    user_api_key = st.text_input(
        "APIキーを入力してください", 
        value=api_key,
        type="password", 
        help="Google AI Studio等で無料で取得できるGemini APIキーを入力します。"
    )
    
    if user_api_key:
        api_key = user_api_key
        try:
            genai.configure(api_key=api_key)
            st.success("🔑 API連携中")
        except Exception as e:
            st.error(f"APIキー設定エラー: {e}")
    else:
        st.warning("⚠️ APIキー未設定です")
        
    st.markdown("---")
    doc_type = st.radio("文末モード:", ["です・ます調（通知表）", "である・した調（要録）"])
    ending_rule = "文末は「です・ます」調で統一。" if "です" in doc_type else "文末は「である・した」調で統一。"
    max_char_limit = st.slider("所見の文字数目安（AI生成）:", min_value=50, max_value=300, value=150, step=10)


# ヘルパー関数
def replace_docx_tags(doc, data_dict):
    def replace_in_paragraphs(paragraphs):
        for p in paragraphs:
            for key, value in data_dict.items():
                tag = f"{{{{{key}}}}}"
                if tag in p.text:
                    p.text = p.text.replace(tag, str(value) if value is not None else "")

    def replace_in_table(table):
        for row in table.rows:
            for cell in row.cells:
                replace_in_paragraphs(cell.paragraphs)
                for nested_table in cell.tables:
                    replace_in_table(nested_table)

    replace_in_paragraphs(doc.paragraphs)
    for table in doc.tables:
        replace_in_table(table)
    for section in doc.sections:
        replace_in_paragraphs(section.header.paragraphs)
        replace_in_paragraphs(section.footer.paragraphs)
        for table in section.header.tables:
            replace_in_table(table)
        for table in section.footer.tables:
            replace_in_table(table)

    return doc

# 匿名化処理を内包した安全なAI生成関数
def safe_generate_content(prompt_text, student_name_map=None):
    if not api_key:
        return "⚠️ Gemini API Keyが設定されていません。サイドバーから設定してください。"
    
    # 仮名化（マスキング）処理
    masked_prompt = prompt_text
    reverse_map = {}
    
    if use_anonymize and student_name_map:
        for idx, real_name in enumerate(student_name_map):
            fake_name = f"生徒{chr(65 + idx)}"  # 生徒A, 生徒B...
            masked_prompt = masked_prompt.replace(real_name, fake_name)
            reverse_map[fake_name] = real_name

    try:
        model = genai.GenerativeModel("gemini-2.5-flash")
        res = model.generate_content(masked_prompt)
        res_text = res.text.strip()

        # 復元処理
        if use_anonymize and reverse_map:
            for fake_name, real_name in reverse_map.items():
                res_text = res_text.replace(fake_name, real_name)

        return res_text
    except Exception as e:
        return f"❌ AI生成中にエラーが発生しました: {str(e)}"

def get_all_classes():
    classes = sorted(st.session_state.student_master["クラス"].dropna().unique().tolist())
    return classes if classes else ["1年1組"]

def get_active_students(cls_name=None):
    df = st.session_state.student_master
    active_df = df[df["ステータス"] == "在籍"]
    if cls_name:
        active_df = active_df[active_df["クラス"] == cls_name]
    return active_df["氏名"].tolist()


selected_menu = st.session_state.selected_menu

# ==========================================
# メイン画面処理
# ==========================================

# ------------------------------------------
# 機能0: ⓪ 担任＆授業担当 名簿管理
# ------------------------------------------
if selected_menu == "⚙️ ⓪ 担任＆授業担当 名簿管理":
    st.subheader("⚙️ 全5クラス（200名）マスター名簿管理")
    m_tab1, m_tab2, m_tab3 = st.tabs(["🏫 担任＆全担当クラス名簿・手動編集", "📥 Excel/CSV ファイルから一括取り込み", "📚 授業担当クラス名簿（コピペ・CSV出力）"])
    
    with m_tab1:
        col_m1, col_m2 = st.columns([1.3, 1])
        with col_m1:
            st.markdown("### 📋 全生徒マスター名簿（200名）")
            edited_master = st.data_editor(
                st.session_state.student_master,
                num_rows="dynamic",
                use_container_width=True,
                height=450,
                key="master_editor"
            )
            st.session_state.student_master = edited_master

        with col_m2:
            st.markdown("### 🔄 転入・転出手続き")
            action_type = st.radio("手続き種別:", ["生徒の新規登録・転入処理", "年度途中 転出（除籍）処理"])
            if action_type == "生徒の新規登録・転入処理":
                with st.form("trans_in_form"):
                    current_classes = get_all_classes()
                    in_cls_select = st.selectbox("登録クラス", current_classes + ["新規クラスを直接入力"])
                    in_cls_custom = st.text_input("新規クラス名（例: 2年3組）") if in_cls_select == "新規クラスを直接入力" else ""
                    target_cls = in_cls_custom if in_cls_select == "新規クラスを直接入力" else in_cls_select
                    in_num = st.number_input("出席番号", min_value=1, max_value=50, value=41)
                    in_name = st.text_input("生徒氏名")
                    in_gender = st.selectbox("性別", ["男", "女"])
                    in_date = st.date_input("登録日")
                    if st.form_submit_button("➕ 生徒を登録する") and in_name.strip() and target_cls:
                        new_st = pd.DataFrame([{"クラス": target_cls, "出席番号": in_num, "氏名": in_name, "性別": in_gender, "ステータス": "在籍", "異動日": str(in_date), "備考": "転入"}])
                        st.session_state.student_master = pd.concat([st.session_state.student_master, new_st], ignore_index=True)
                        st.success(f"{target_cls} に {in_name} さんを登録しました！")
                        st.rerun()

            else:
                with st.form("trans_out_form"):
                    active_list = get_active_students()
                    out_name = st.selectbox("転出生徒を選択", active_list)
                    out_date = st.date_input("転出日")
                    out_reason = st.text_input("転出理由")
                    if st.form_submit_button("⚠️ 転出処理を実行"):
                        idx = st.session_state.student_master[st.session_state.student_master["氏名"] == out_name].index
                        st.session_state.student_master.loc[idx, "ステータス"] = "転出"
                        st.session_state.student_master.loc[idx, "異動日"] = str(out_date)
                        st.session_state.student_master.loc[idx, "備考"] = out_reason
                        st.warning(f"{out_name} さんの転出処理を完了しました。")
                        st.rerun()

    with m_tab2:
        st.markdown("### 📥 既存の名簿ファイル（.xlsx / .csv）を一括取り込み")
        uploaded_file = st.file_uploader("名簿ファイルをアップロード:", type=["csv", "xlsx"])
        if uploaded_file is not None:
            try:
                imp_df = pd.read_csv(uploaded_file) if uploaded_file.name.endswith(".csv") else pd.read_excel(uploaded_file)
                st.write("📖 **プレビュー:**")
                st.dataframe(imp_df.head(), use_container_width=True)
                req_cols = ["クラス", "出席番号", "氏名"]
                if all(col in imp_df.columns for col in req_cols):
                    if "性別" not in imp_df.columns: imp_df["性別"] = "未設定"
                    if "ステータス" not in imp_df.columns: imp_df["ステータス"] = "在籍"
                    if "異動日" not in imp_df.columns: imp_df["異動日"] = ""
                    if "備考" not in imp_df.columns: imp_df["備考"] = "ファイル取込"
                    if st.button("🚀 この名簿データをシステムに取り込む"):
                        st.session_state.student_master = imp_df[st.session_state.student_master.columns]
                        st.success("🎉 名簿データベースの更新が完了しました！")
                        st.rerun()
                else:
                    st.error("必須列（クラス, 出席番号, 氏名）が含まれていません。")
            except Exception as e:
                st.error(f"読み込みエラー: {e}")

    with m_tab3:
        st.markdown("### 📚 担当5クラスの個別名簿抽出")
        all_classes = get_all_classes()
        selected_teach_cls = st.multiselect("抽出対象クラス選択:", all_classes, default=all_classes)
        if selected_teach_cls:
            sub_df = st.session_state.student_master[
                (st.session_state.student_master["クラス"].isin(selected_teach_cls)) &
                (st.session_state.student_master["ステータス"] == "在籍")
            ][["クラス", "出席番号", "氏名", "性別"]].sort_values(by=["クラス", "出席番号"])
            st.dataframe(sub_df, use_container_width=True, height=300)
            st.text_area("Excel貼り付け用タブ区切りテキスト (Ctrl+V用):", sub_df.to_csv(sep='\t', index=False), height=100)

# ------------------------------------------
# 機能1: ① 日々メモ・クイックスタンプ & 所見
# ------------------------------------------
elif selected_menu == "📝 ① 日々メモ・クイックスタンプ & 所見":
    st.subheader("📝 日々の観察記録 ＆ クイックスタンプ登録 ＆ AI所見生成")
    col_a, col_b = st.columns([1.1, 1.1])
    all_cls = get_all_classes()
    
    with col_a:
        st.markdown("### 📌 メモの追加登録")
        f_class = st.selectbox("クラス", all_cls, index=0, key="quick_cls")
        c_students = get_active_students(f_class)
        f_name = st.selectbox("生徒氏名", c_students, key="quick_st")
        f_date = st.date_input("日付", key="quick_date")
        f_cat = st.selectbox("対象分野", ["数学", "総合・行動の記録", "国語・他", "特別活動", "生活指導"], key="quick_cat")

        st.markdown("**⚡ 1タップ クイックスタンプ登録:**")
        st_col1, st_col2, st_col3 = st.columns(3)
        
        stamps = [
            ("🙋‍♂️ 挙手・発言意欲◎", "数学", "授業中に積極的に挙手し発言することができた。"),
            ("💡 自力解決できた", "数学", "難しい問題に対して粘り強く思考し自力で正解を導き出した。"),
            ("🤝 班活動をリード", "総合・行動", "グループワークで意見をまとめ、周囲をよくサポートしていた。"),
            ("🧹 清掃・手伝い貢献", "総合・行動", "自分の担当領域が終わった後も進んで周囲の手伝いを行った。"),
            ("📖 朝読書・集中", "特別活動", "静かに集中して朝の読書活動に取り組めた。"),
            ("⚠️ 提出物の遅れ", "生活指導", "提出物の期限について個別に声かけを行った。")
        ]

        def add_stamp_log(cat, memo):
            new_row = pd.DataFrame([{"日付": str(f_date), "クラス": f_class, "出席番号": 1, "氏名": f_name, "対象分野": cat, "観察メモ": memo}])
            st.session_state.daily_logs = pd.concat([st.session_state.daily_logs, new_row], ignore_index=True)
            st.toast(f"✅ {f_name} さんにスタンプ『{memo[:10]}...』を記録しました！")

        for idx, (label, cat, text) in enumerate(stamps):
            target_col = [st_col1, st_col2, st_col3][idx % 3]
            if target_col.button(label, key=f"stamp_{idx}"):
                add_stamp_log(cat, text)

        st.markdown("---")
        st.markdown("**✍️ テキスト手動入力**")
        with st.form("add_log_form", clear_on_submit=True):
            f_memo = st.text_area("自由記述観察メモ", placeholder="授業や学級活動での具体的な様子...")
            if st.form_submit_button("📥 詳細メモを保存") and f_name and f_memo:
                new_row = pd.DataFrame([{"日付": str(f_date), "クラス": f_class, "出席番号": 1, "氏名": f_name, "対象分野": f_cat, "観察メモ": f_memo}])
                st.session_state.daily_logs = pd.concat([st.session_state.daily_logs, new_row], ignore_index=True)
                st.success(f"{f_name} さんのメモを追加しました！")

    with col_b:
        st.markdown("### ✨ 蓄積メモからAI所見生成")
        selected_student = f_name
        
        student_memos = st.session_state.daily_logs[st.session_state.daily_logs["氏名"] == selected_student]
        st.write(f"📜 **{selected_student} さんの蓄積メモ（{len(student_memos)}件）:**")
        st.dataframe(student_memos[["日付", "対象分野", "観察メモ"]], use_container_width=True, height=180)
        
        if st.button("🪄 蓄積メモから所見文案を自動生成", type="primary"):
            if not student_memos.empty:
                combined_memos = "\n".join(student_memos["観察メモ"].tolist())
                prompt = f"生徒『{selected_student}』の蓄積メモ:\n{combined_memos}\n\n上記メモをもとに通知表用の所見文案を作成してください。文字数は約{max_char_limit}文字程度。{ending_rule}"
                with st.spinner("AIが個人情報を保護しながら所見文案を作成中..."):
                    generated_text = safe_generate_content(prompt, student_name_map=[selected_student])
                    st.text_area("生成された所見文案:", value=generated_text, height=160)
            else:
                st.warning("この生徒の観察メモがまだ登録されていません。")

# ------------------------------------------
# 機能2: ② 学級日常ダッシュボード（観察アラート）
# ------------------------------------------
elif selected_menu == "🔔 ② 学級日常ダッシュボード（観察アラート）":
    st.subheader("🔔 学級日常ダッシュボード（観察メモ不足・見守りアラート）")
    st.caption("特定生徒への観察の偏りを防ぎ、過去14日間観察記録がない生徒を自動抽出します。")

    target_dash_cls = st.selectbox("対象クラス選択:", get_all_classes(), key="dash_cls_sel")
    cls_students = get_active_students(target_dash_cls)
    
    logs_df = st.session_state.daily_logs[st.session_state.daily_logs["クラス"] == target_dash_cls].copy()
    logs_df["日付_dt"] = pd.to_datetime(logs_df["日付"])
    
    two_weeks_ago = datetime.now() - timedelta(days=14)
    
    summary_data = []
    alert_students = []

    for name in cls_students:
        st_logs = logs_df[logs_df["氏名"] == name]
        total_count = len(st_logs)
        recent_count = len(st_logs[st_logs["日付_dt"] >= two_weeks_ago])
        last_date = st_logs["日付"].max() if not st_logs.empty else "記録なし"
        
        is_alert = recent_count == 0
        if is_alert:
            alert_students.append(name)

        summary_data.append({
            "氏名": name,
            "累計メモ数": total_count,
            "直近14日間のメモ": f"{recent_count} 件",
            "最終記録日": last_date,
            "状態アラート": "🚨 14日以上メモなし" if is_alert else "✅ 順調に蓄積中"
        })

    sum_df = pd.DataFrame(summary_data)

    col_d1, col_d2 = st.columns([1, 2])
    with col_d1:
        st.metric("クラス総人数", f"{len(cls_students)} 名")
        st.metric("🚨 観察メモ不足（要声かけ）生徒", f"{len(alert_students)} 名")

        if alert_students:
            st.markdown("<div class='alert-card'><b>💡 以下の生徒は最近記録がありません:</b><br>" + "、".join(alert_students) + "</div>", unsafe_allow_html=True)

    with col_d2:
        st.markdown("### 📋 クラス全員の記録蓄積状況")
        st.dataframe(sum_df, use_container_width=True, height=350)

# ------------------------------------------
# 機能3: ③ CSV一括生成
# ------------------------------------------
elif selected_menu == "📁 ③ CSV一括生成":
    st.subheader("📁 クラス全員の蓄積メモから一括所見生成")
    target_cls = st.selectbox("一括生成対象クラス:", get_all_classes(), key="t2_cls")
    
    if st.button("🚀 クラス全員の所見をAI一括生成"):
        cls_memos = st.session_state.daily_logs[st.session_state.daily_logs["クラス"] == target_cls]
        results = []
        progress_bar = st.progress(0)
        unique_names = get_active_students(target_cls)
        
        for i, name in enumerate(unique_names):
            group = cls_memos[cls_memos["氏名"] == name]
            all_memos = " / ".join(group["観察メモ"].tolist()) if not group.empty else "日々の授業に真面目に取り組んでいる。"
            prompt = f"生徒:{name} メモ:{all_memos} の通知表所見を作成。文字数は約{max_char_limit}文字。{ending_rule}"
            
            gen_text = safe_generate_content(prompt, student_name_map=[name])
            results.append({"氏名": name, "まとめメモ": all_memos, "生成所見": gen_text})
            progress_bar.progress((i + 1) / len(unique_names))
                
        res_df = pd.DataFrame(results)
        st.dataframe(res_df, use_container_width=True, height=300)
        csv_data = res_df.to_csv(index=False).encode('utf-8-sig')
        st.download_button(f"📥 {target_cls} 全員の所見一括CSVをダウンロード", csv_data, f"{target_cls}_一括所見データ.csv", "text/csv")

# ------------------------------------------
# 機能4: ④ 所見自動校正
# ------------------------------------------
elif selected_menu == "🔍 ④ 所見データ自動校正":
    st.subheader("🔍 所見データの自動校正 & 不適切表現チェック")
    c1_students = get_active_students("1年1組")
    student_for_check = st.selectbox("校正を試す生徒を選択:", c1_students, key="chk_st")
    
    memos_text = " ".join(st.session_state.daily_logs[st.session_state.daily_logs["氏名"] == student_for_check]["観察メモ"].tolist())
    sample_text = st.text_area("校正対象テキスト:", value=f"{student_for_check}さんは、" + memos_text, height=120)
        
    if st.button("🛡️ AI誤字脱字・表現校正を実行", type="primary"):
        if sample_text:
            prompt = f"誤字脱字チェックおよび保護者目線での適切な文章校正を行ってください:\n{sample_text}\n{ending_rule}"
            res_text = safe_generate_content(prompt, student_name_map=[student_for_check])
            st.markdown("### 💡 校正結果アドバイス:")
            st.info(res_text)

# ------------------------------------------
# 機能5: ⑤ 蓄積連動カルテ
# ------------------------------------------
elif selected_menu == "💬 ⑤ 蓄積連動カルテ":
    st.subheader("💬 保護者面談用カルテ（蓄積メモ＋テスト成績の自動連携）")
    kart_student = st.selectbox("面談対象生徒を選択:", get_active_students("1年1組"), key="kart_st")
    
    st_memos = st.session_state.daily_logs[st.session_state.daily_logs["氏名"] == kart_student]
    st_scores = st.session_state.subject_scores[st.session_state.subject_scores["氏名"] == kart_student]
    
    col_k1, col_k2 = st.columns(2)
    with col_k1:
        st.markdown(f"### 📜 日々の観察記録（{len(st_memos)}件）")
        st.dataframe(st_memos[["日付", "対象分野", "観察メモ"]], use_container_width=True, height=200)
    with col_k2:
        st.markdown("### 📊 自教科テスト＆観点成績")
        st.dataframe(st_scores[["中間テスト", "期末テスト", "見込み点", "観点1_評価", "観点2_評価", "観点3_評価", "★確定評定"]], use_container_width=True)
        
    if st.button("📋 面談用トークポイントカルテをAI生成", type="primary"):
        memo_concat = " ".join(st_memos['観察メモ'].tolist())
        score_info = st_scores.to_dict(orient="records")[0] if not st_scores.empty else {}
        prompt = f"生徒『{kart_student}』の観察記録:{memo_concat}\nテスト成績:中間{score_info.get('中間テスト')}点, 期末{score_info.get('期末テスト')}点, 最終評定{score_info.get('★確定評定')}\n面談で保護者に伝える【1.学習面・生活面の成長点 2.今後の課題 3.家庭での連携アドバイス】を簡潔に作成してください。"
        res_text = safe_generate_content(prompt, student_name_map=[kart_student])
        st.info("💡 **AI生成 面談用カルテシート:**")
        st.markdown(res_text)

# ------------------------------------------
# 機能6: ⑥ 成績・観点A/B/C算出＆人間調整
# ------------------------------------------
elif selected_menu == "📊 ⑥ 成績・観点A/B/C算出＆人間調整":
    st.subheader("📊 成績・観点A/B/C評価算出 ＆ カッティングポイント設定 ＆ 人間調整ワークスペース")
    
    sel_eval_cls = st.selectbox("対象クラス切替:", get_all_classes(), key="t5_cls")

    with st.expander("⚙️ 【設定】観点別評価（A/B/C）のカッティングポイント（しきい値）設定", expanded=False):
        st.caption("各観点の点数（100点満点換算）をどのラインでA・B・C評価に自動換算するかを設定できます。")
        col_cp1, col_cp2, col_cp3 = st.columns(3)
        with col_cp1:
            st.markdown("**観点1 (知識・技能)**")
            cut_k1_a = st.number_input("A評価のボーダー (点以上)", value=80, key="cp_k1_a")
            cut_k1_b = st.number_input("B評価のボーダー (点以上)", value=50, key="cp_k1_b")
        with col_cp2:
            st.markdown("**観点2 (思考・判断・表現)**")
            cut_k2_a = st.number_input("A評価のボーダー (点以上)", value=80, key="cp_k2_a")
            cut_k2_b = st.number_input("B評価のボーダー (点以上)", value=50, key="cp_k2_b")
        with col_cp3:
            st.markdown("**観点3 (主体的に学習に取り組む態度)**")
            cut_k3_a = st.number_input("A評価のボーダー (点以上)", value=80, key="cp_k3_a")
            cut_k3_b = st.number_input("B評価のボーダー (点以上)", value=50, key="cp_k3_b")

    st.markdown("---")

    cls_score_df = st.session_state.subject_scores[st.session_state.subject_scores["クラス"] == sel_eval_cls].copy()

    st.markdown("### ✏️ 成績＆評価データ・人間微調整シート")
    st.caption("💡 **調整方法:** 右側の「★確定評定」列を必要に応じて打ち替えてください。自動評定と異なる場合は「⚠️ 変更済」フラグが自動的に立ちます。")

    columns_order = [
        "出席番号", "氏名", "中間テスト", "期末テスト", "見込み点",
        "観点1_知識(点)", "観点1_評価", 
        "観点2_思考(点)", "観点2_評価", 
        "観点3_主体性(点)", "観点3_評価", 
        "自動評定", "★確定評定", "調整フラグ", "調整理由"
    ]
    
    for col in columns_order:
        if col not in cls_score_df.columns:
            cls_score_df[col] = ""

    display_df = cls_score_df[columns_order]

    edited_scores = st.data_editor(
        display_df,
        num_rows="dynamic",
        use_container_width=True,
        height=400,
        key=f"score_editor_{sel_eval_cls}"
    )

    col_btn1, col_btn2 = st.columns([1.2, 1])
    with col_btn1:
        if st.button("⚡ 設定したカッティングポイントで A/B/C & 自動評定を再計算"):
            def recalculate_row(row):
                k1 = row["観点1_知識(点)"]
                row["観点1_評価"] = "A" if k1 >= cut_k1_a else ("B" if k1 >= cut_k1_b else "C")
                
                k2 = row["観点2_思考(点)"]
                row["観点2_評価"] = "A" if k2 >= cut_k2_a else ("B" if k2 >= cut_k2_b else "C")
                
                k3 = row["観点3_主体性(点)"]
                row["観点3_評価"] = "A" if k3 >= cut_k3_a else ("B" if k3 >= cut_k3_b else "C")

                avg = (k1 + k2 + k3) / 3
                auto_g = 5 if avg >= 85 else (4 if avg >= 70 else (3 if avg >= 55 else (2 if avg >= 40 else 1)))
                row["自動評定"] = auto_g

                if pd.isna(row["★確定評定"]) or str(row["★確定評定"]).strip() == "":
                    row["★確定評定"] = auto_g

                if str(row["自動評定"]) != str(row["★確定評定"]):
                    row["調整フラグ"] = "⚠️ 変更済"
                else:
                    row["調整フラグ"] = "―"
                
                return row

            updated_df = edited_scores.apply(recalculate_row, axis=1)
            st.session_state.subject_scores.update(updated_df)
            st.success("🎉 指定したカッティングポイントに基づき、A/B/C評価と評定・調整フラグを最新化しました！")
            st.rerun()

    with col_btn2:
        if st.button("💾 人間調整後の評定データをデータベースに確定保存"):
            def check_flag(row):
                if str(row["自動評定"]) != str(row["★確定評定"]):
                    row["調整フラグ"] = "⚠️ 変更済"
                else:
                    row["調整フラグ"] = "―"
                return row
            
            final_df = edited_scores.apply(check_flag, axis=1)
            st.session_state.subject_scores.update(final_df)
            st.success("✅ 人間調整後の確定評定データをデータベースに保存しました！")

# ------------------------------------------
# 機能7: ⑦ 学期推移ダッシュボード
# ------------------------------------------
elif selected_menu == "📈 ⑦ 学期推移ダッシュボード":
    st.subheader("📈 学期・テスト別 成績推移ダッシュボード")
    dash_cls = st.selectbox("ダッシュボード対象クラス:", get_all_classes(), key="t6_cls")
    
    cls_scores = st.session_state.subject_scores[st.session_state.subject_scores["クラス"] == dash_cls]
    
    st.markdown(f"### 📊 {dash_cls} 中間テスト vs 期末テスト 推移グラフ")
    chart_data = cls_scores.set_index("氏名")[["中間テスト", "期末テスト"]]
    st.line_chart(chart_data)
    
    st.markdown("### 📋 成績下降・フォロー対象者（期末テストで点数が下がった生徒）")
    down_students = cls_scores[cls_scores["期末テスト"] < cls_scores["中間テスト"]]
    st.dataframe(down_students[["出席番号", "氏名", "中間テスト", "期末テスト", "★確定評定", "調整フラグ"]], use_container_width=True)

# ------------------------------------------
# 機能8: ⑧ 担任用 全教科成績集約
# ------------------------------------------
elif selected_menu == "🔄 ⑧ 担任用 全教科成績集約":
    st.subheader("🔄 各教科の成績ファイル自動名寄せ統合（ダミーデータ機能付き）")
    st.caption("他教科の担任から集まった個別CSVファイルを1人1行の全教科シートに合体します。")
    
    if st.button("🎲 他教科（国語・英語・理科・社会）の集約用ダミーデータを自動生成して結合テスト"):
        c1_students = st.session_state.student_master[st.session_state.student_master["クラス"] == "1年1組"]["氏名"].tolist()
        
        kokugo_df = pd.DataFrame([{"氏名": n, "国語_評定": random.randint(2, 5), "国語_観点_知識": random.randint(60, 95)} for n in c1_students])
        eigo_df = pd.DataFrame([{"氏名": n, "英語_評定": random.randint(2, 5), "英語_観点_知識": random.randint(55, 98)} for n in c1_students])
        
        math_df = st.session_state.subject_scores[st.session_state.subject_scores["クラス"] == "1年1組"][["氏名", "★確定評定"]].rename(columns={"★確定評定": "数学_評定"})
        
        merged_all = pd.merge(math_df, kokugo_df, on="氏名")
        merged_all = pd.merge(merged_all, eigo_df, on="氏名")
        
        st.success("🎉 1年1組の数学・国語・英語の成績データを『氏名』で完璧に合体しました！")
        st.dataframe(merged_all, use_container_width=True, height=350)
        
        csv_m = merged_all.to_csv(index=False).encode('utf-8-sig')
        st.download_button("📥 1年1組 全教科統合成績シート（CSV）を保存", csv_m, "1年1組_全教科成績統合表.csv", "text/csv")

# ------------------------------------------
# 機能9: 🖨️ 完成版プレビュー ＆ ファイルダウンロード（通知表・指導要録）
# ------------------------------------------
elif selected_menu == "🖨️ ⑨ 完成版プレビュー ＆ ファイルダウンロード（通知表・指導要録）":
    st.subheader("🖨️ 完成版プレビュー ＆ ファイルダウンロード（通知表・指導要録）")
    st.caption("通知表および指導要録のデータを差し込み表示・各種形式（Word/PDF）で出力できます。")
    
    col_out1, col_out2 = st.columns([1.1, 1.9])
    
    with col_out1:
        st.markdown("### 1. 出力種別・対象生徒の選択")
        doc_category = st.radio("📄 出力対象文書:", ["通知表", "指導要録"], key="doc_cat_select")
        
        out_cls = st.selectbox("クラス選択:", get_all_classes(), index=0, key="t8_cls")
        out_students = get_active_students(out_cls)
        print_student = st.selectbox("生徒選択:", out_students)
        
        st.markdown("---")
        st.markdown("### 2. 枠組み（テンプレート）取り込み")
        st.caption("※Word (.docx) または PDF (.pdf) のテンプレート枠組みを取り込めます。")
        template_file = st.file_uploader(
            f"学校独自{doc_category}テンプレート (.docx / .pdf)", 
            type=["docx", "pdf"],
            key="template_file_upload"
        )
        
        st.markdown("---")
        st.markdown("### 3. ダウンロード設定")
        output_format = st.selectbox("出力フォーマット:", ["Word形式 (.docx)", "PDF形式 (.pdf)"], key="out_format_select")

    with col_out2:
        st.markdown(f"### 📄 差し込みプレビュー（{doc_category}：{out_cls} {print_student} 様）")
        
        if print_student:
            st_info = st.session_state.student_master[st.session_state.student_master["氏名"] == print_student].iloc[0].to_dict()
            score_match = st.session_state.subject_scores[st.session_state.subject_scores["氏名"] == print_student]
            st_score = score_match.iloc[0].to_dict() if not score_match.empty else {}
            
            mid_display = str(st_score.get('中間テスト', '-'))
            if pd.isna(st_score.get('中間テスト')) or st_score.get('中間テスト') is None:
                if st_score.get('見込み点'):
                    mid_display = f"未受検 (見込み点: {st_score.get('見込み点')}点)"
                else:
                    mid_display = "未受検"

            default_remark = st_score.get('総合所見', '（所見データ準備完了）')
            if doc_category == "指導要録":
                default_remark = default_remark.replace("でした。", "であった。").replace("ました。", "した。").replace("です。", "である。")

            st.markdown(f"""
            <div class="student-card">
                <div style="display:flex; justify-content:space-between; align-items:center;">
                    <h3 style="margin:0;">【{st.session_state.school_year}年度 1学期 {doc_category}プレビュー】</h3>
                    <span style="background-color:#007bff; color:white; padding:3px 8px; border-radius:4px; font-size:0.8em;">
                        取り込み形式: {template_file.name if template_file else '標準システムフォーマット'}
                    </span>
                </div>
                <hr>
                <p><strong>所属:</strong> {st_info.get('クラス')} | <strong>出席番号:</strong> {st_info.get('出席番号')}番</p>
                <p><strong>氏名:</strong> <span style="font-size:1.3em; font-weight:bold;">{print_student}</span></p>
                <p><strong>学級担任:</strong> {st.session_state.teacher_name}</p>
                <hr>
                <h4>📊 学習評価・記録成績（差し込み完了）</h4>
                <ul>
                    <li>中間テスト: <strong>{mid_display}</strong> / 期末テスト: <strong>{st_score.get('期末テスト', '-')}</strong> 点</li>
                    <li>観点1 (知識・技能): <strong>{st_score.get('観点1_評価', '-')}</strong> ({st_score.get('観点1_知識(点)', '-')}点)</li>
                    <li>観点2 (思考・判断・表現): <strong>{st_score.get('観点2_評価', '-')}</strong> ({st_score.get('観点2_思考(点)', '-')}点)</li>
                    <li>観点3 (主体的に取り組む態度): <strong>{st_score.get('観点3_評価', '-')}</strong> ({st_score.get('観点3_主体性(点)', '-')}点)</li>
                    <li>確定学習評定（5段階）: <strong style="font-size:1.4em; color:#d9534f;">{st_score.get('★確定評定', '-')}</strong> (自動算出: {st_score.get('自動評定', '-')}) {st_score.get('調整フラグ', '')}</li>
                </ul>
                <hr>
                <h4>📝 {doc_category} 所見欄（差し込み完了）</h4>
                <p style="background-color: white; padding: 12px; border-radius: 6px; border: 1px solid #ccc; line-height: 1.6;">
                    {default_remark}
                </p>
            </div>
            """, unsafe_allow_html=True)
            
            doc_data = {
                "年度": st.session_state.school_year,
                "担任名": st.session_state.teacher_name,
                "クラス": st_info.get('クラス'),
                "出席番号": st_info.get('出席番号'),
                "氏名": print_student,
                "中間": mid_display,
                "期末": st_score.get('期末テスト', ''),
                "観点1": st_score.get('観点1_評価', ''),
                "観点2": st_score.get('観点2_評価', ''),
                "観点3": st_score.get('観点3_評価', ''),
                "評定": st_score.get('★確定評定', ''),
                "総合所見": default_remark
            }

            download_label = f"📥 {print_student} さんの {doc_category}（{output_format}）を出力・ダウンロード"
            
            if "Word" in output_format:
                try:
                    if template_file and template_file.name.endswith(".docx"):
                        doc = Document(template_file)
                    else:
                        doc = Document()
                        doc.add_heading(f"{doc_category} - {print_student} 様", 0)
                        doc.add_paragraph(f"年度: {st.session_state.school_year}年度 | 担任: {st.session_state.teacher_name}")
                        doc.add_paragraph(f"クラス: {st_info.get('クラス')}  出席番号: {st_info.get('出席番号')}")
                        doc.add_paragraph(f"中間テスト: {mid_display} | 期末テスト: {st_score.get('期末テスト', '')}点")
                        doc.add_paragraph(f"数学評定: {st_score.get('★確定評定', '')}")
                        doc.add_paragraph(f"{doc_category}所見:\n{default_remark}")
                    
                    filled_doc = replace_docx_tags(doc, doc_data)
                    out_buffer = io.BytesIO()
                    filled_doc.save(out_buffer)
                    
                    st.download_button(
                        label=download_label,
                        data=out_buffer.getvalue(),
                        file_name=f"{doc_category}_{print_student}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="primary"
                    )
                except Exception as e:
                    st.error(f"Wordファイルの生成中にエラーが発生しました: {e}")

            elif "PDF" in output_format:
                pdf_text = f"""==================================================
【{st.session_state.school_year}年度 1学期 {doc_category}】
==================================================
■ 児童生徒名 : {print_student}
■ 所属      : {st_info.get('クラス')}  出席番号: {st_info.get('出席番号')}番
■ 学級担任  : {st.session_state.teacher_name}
--------------------------------------------------
【学習の記録】
・中間テスト : {mid_display}
・期末テスト : {st_score.get('期末テスト', '')} 点
・観点1(知識): {st_score.get('観点1_評価', '')}
・観点2(思考): {st_score.get('観点2_評価', '')}
・観点3(主体): {st_score.get('観点3_評価', '')}
・確定評定   : {st_score.get('★確定評定', '')}
--------------------------------------------------
【{doc_category} 所見】
{default_remark}
==================================================
"""
                pdf_buffer = io.BytesIO(pdf_text.encode("utf-8"))
                st.download_button(
                    label=download_label,
                    data=pdf_buffer.getvalue(),
                    file_name=f"{doc_category}_{print_student}.pdf",
                    mime="application/pdf",
                    type="primary"
                )